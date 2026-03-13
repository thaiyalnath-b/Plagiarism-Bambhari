# # import os
# # import pdfplumber
# # import docx

# # # OCR imports (for scanned thesis PDFs)
# # import pytesseract
# # from PIL import Image



# # # --------------------------------------------------
# # # OPTIONAL: SET TESSERACT PATH (Windows only)
# # # Change path if different in your system
# # # --------------------------------------------------

# # pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# # # "C:\Program Files\Tesseract-OCR\tessdata"

# # # --------------------------------------------------
# # # MAIN TEXT EXTRACTION FUNCTION
# # # --------------------------------------------------
# # def extract_text(file_path):
# #     """
# #     Extract text safely from supported files.

# #     Supported:
# #         - PDF (normal + scanned OCR)
# #         - DOCX
# #         - TXT

# #     Always returns STRING (never None)
# #     """

# #     if not file_path or not os.path.exists(file_path):
# #         return ""

# #     file_path = file_path.lower()

# #     try:
# #         # ---------------- PDF ----------------
# #         if file_path.endswith(".pdf"):
# #             return extract_pdf_text(file_path)

# #         # ---------------- DOCX ----------------
# #         elif file_path.endswith(".docx"):
# #             return extract_docx_text(file_path)

# #         # ---------------- TXT ----------------
# #         elif file_path.endswith(".txt"):
# #             return extract_txt_text(file_path)

# #     except Exception as e:
# #         print("TEXT EXTRACTION ERROR:", e)

# #     return ""


# # # --------------------------------------------------
# # # PDF EXTRACTION (NORMAL + OCR FALLBACK)
# # # --------------------------------------------------
# # def extract_pdf_text(file_path):

# #     text = ""

# #     try:
# #         with pdfplumber.open(file_path) as pdf:
# #             MAX_PAGES = 12
# #             for page in pdf.pages[:MAX_PAGES]:
# #                 page_text = page.extract_text()

# #                 if page_text:
# #                     text += page_text + "\n"

# #             # --------------------------------------------------
# #             # OCR FALLBACK (for scanned/image thesis PDFs)
# #             # --------------------------------------------------
# #             if len(text.strip()) < 50:
# #                 print("⚠ OCR MODE ACTIVATED (Scanned PDF detected)")
                
# #                 for page in pdf.pages:
# #                     try:
# #                         # Convert page to image
# #                         img = page.to_image(resolution=300).original

# #                         ocr_text = pytesseract.image_to_string(img)

# #                         if ocr_text:
# #                             text += ocr_text + "\n"

# #                     except Exception as ocr_error:
# #                         print("OCR page failed:", ocr_error)

# #     except Exception as e:
# #         print("PDF extraction failed:", e)

# #     return text.strip()


# # # --------------------------------------------------
# # # DOCX EXTRACTION
# # # --------------------------------------------------
# # def extract_docx_text(file_path):

# #     text = ""

# #     try:
# #         document = docx.Document(file_path)

# #         for paragraph in document.paragraphs:
# #             if paragraph.text:
# #                 text += paragraph.text + "\n"

# #     except Exception as e:
# #         print("DOCX extraction failed:", e)

# #     return text.strip()


# # # --------------------------------------------------
# # # TXT EXTRACTION
# # # --------------------------------------------------
# # def extract_txt_text(file_path):

# #     try:
# #         with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
# #             return f.read().strip()

# #     except Exception as e:
# #         print("TXT extraction failed:", e)

# #     return ""


# import pdfplumber
# import docx
# import os

# RENDER_ENV = os.environ.get("RENDER") is not None

# # OPTIONAL OCR
# try:
#     import pytesseract
#     from pdf2image import convert_from_path
#     OCR_AVAILABLE = True
# except:
#     OCR_AVAILABLE = False


# # ---------------------------------------
# # MAIN ENTRY
# # ---------------------------------------
# def extract_text(file_path):

#     if not os.path.exists(file_path):
#         print("file not found")
#         return ""

#     file_path = file_path.lower()

#     if file_path.endswith(".pdf"):
#         return extract_pdf_fast(file_path)

#     elif file_path.endswith(".docx"):
#         doc = docx.Document(file_path)
#         return "\n".join(p.text for p in doc.paragraphs)

#     elif file_path.endswith(".txt"):
#         with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
#             return f.read()

#     return ""


# # ---------------------------------------
# # FAST PDF EXTRACTION (SMART)
# # ---------------------------------------
# def extract_pdf_fast(file_path):

#     text = ""
#     pages_checked = 0

#     try:
#         with pdfplumber.open(file_path) as pdf:

#             total_pages = len(pdf.pages)

#             #  SAMPLE ONLY FIRST + MIDDLE + LAST pages
#             sample_indexes = set([
#                 0,
#                 total_pages // 2,
#                 total_pages - 1
#             ])

#             for i in sample_indexes:
#                 if i < 0 or i >= total_pages:
#                     continue

#                 page_text = pdf.pages[i].extract_text()

#                 if page_text:
#                     text += page_text + "\n"

#                 pages_checked += 1

#     except Exception as e:
#         print("PDF TEXT ERROR:", e)

#     #  if enough text found → skip OCR completely
#     if len(text) > 1500:
#         print(" TEXT PDF detected — OCR skipped")
#         return text[:8000]

#     # ---------------------------------------
#     # OCR FALLBACK (LIMITED)
#     # ---------------------------------------
#     if OCR_AVAILABLE:
#         print("⚠ OCR MODE ACTIVATED (limited pages)")

#         try:
#             images = convert_from_path(
#                 file_path,
#                 first_page=1,
#                 last_page=3  # ⭐ OCR ONLY FIRST 3 PAGES
#             )

#             for img in images:
#                 text += pytesseract.image_to_string(img)

#         except Exception as e:
#             print("OCR failed:", e)

#     return text[:8000]

import os
import pdfplumber
import docx

# ======================================
# MAIN ENTRY
# ======================================
def extract_text(file_path):
    """
    Extract text from various file types
    Supports: PDF, DOCX, TXT
    """
    if not os.path.exists(file_path):
        print(" File not found:", file_path)
        return ""

    file_path_lower = file_path.lower()

    if file_path_lower.endswith(".pdf"):
        return extract_pdf_fast(file_path)

    elif file_path_lower.endswith(".docx"):
        return extract_docx(file_path)

    elif file_path_lower.endswith(".txt"):
        return extract_txt(file_path)

    return ""


# ======================================
# EXTRACT FROM TXT
# ======================================
def extract_txt(file_path):
    """Extract text from TXT file"""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        print(f" TXT extracted: {len(text)} characters")
        return text
    except Exception as e:
        print(f" TXT extraction error: {e}")
        return ""


# ======================================
# EXTRACT FROM DOCX
# ======================================
def extract_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        print(f" DOCX extracted: {len(text)} characters")
        return text
    except Exception as e:
        print(f" DOCX extraction error: {e}")
        return ""


# ======================================
# FAST PDF EXTRACTION (RENDER SAFE)
# ======================================
def extract_pdf_fast(file_path):
    """
    Extract text from PDF file
    Samples first page, middle page, and last page for speed
    """
    text = ""

    try:
        with pdfplumber.open(file_path) as pdf:

            total_pages = len(pdf.pages)
            print("📄 PDF Pages:", total_pages)

            if total_pages == 0:
                return ""

            # Determine which pages to sample
            pages_to_extract = set()
            
            # Always include first page
            pages_to_extract.add(0)
            
            # Include last page if more than 1 page
            if total_pages > 1:
                pages_to_extract.add(total_pages - 1)
            
            # Include middle page if more than 2 pages
            if total_pages > 2:
                pages_to_extract.add(total_pages // 2)
            
            # If only a few pages, extract all
            if total_pages <= 5:
                pages_to_extract = set(range(total_pages))

            print(f"📄 Extracting pages: {sorted(pages_to_extract)}")

            for page_num in sorted(pages_to_extract):
                page = pdf.pages[page_num]
                page_text = page.extract_text()
                
                if page_text:
                    text += page_text + "\n\n"

    except Exception as e:
        print(" PDF extraction error:", e)
        import traceback
        traceback.print_exc()
        return ""

    # Clean up text
    text = ' '.join(text.split())
    
    print(f" Extracted characters: {len(text)}")
    
    # Limit to first 8000 chars for performance
    if len(text) > 8000:
        text = text[:8000]
        print(f"📏 Truncated to 8000 chars")

    return text


# ======================================
# EXTRACT FROM UPLOADED FILE (DIRECT)
# ======================================
def extract_from_uploaded_file(uploaded_file):
    """
    Extract text directly from an uploaded file object
    Saves temporarily then extracts
    """
    import tempfile
    
    # Create temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
        for chunk in uploaded_file.chunks():
            tmp_file.write(chunk)
        tmp_path = tmp_file.name
    
    try:
        # Extract text from temporary file
        text = extract_text(tmp_path)
        return text
    finally:
        # Clean up temporary file
        try:
            os.unlink(tmp_path)
        except:
            pass

# import pdfplumber
# import docx
# import os


# # ======================================
# # MAIN ENTRY
# # ======================================
# def extract_text(file_path):

#     if not os.path.exists(file_path):
#         print(" File not found")
#         return ""

#     file_path = file_path.lower()

#     if file_path.endswith(".pdf"):
#         return extract_pdf_text(file_path)

#     elif file_path.endswith(".docx"):
#         return extract_docx_text(file_path)

#     elif file_path.endswith(".txt"):
#         return extract_txt_text(file_path)

#     return ""


# # ======================================
# # FAST PDF EXTRACTION (NO OCR)
# # ======================================
# def extract_pdf_text(file_path):

#     text = ""

#     try:
#         with pdfplumber.open(file_path) as pdf:

#             total_pages = len(pdf.pages)

#             # ⭐ SPEED OPTIMIZATION
#             # only sample pages
#             sample_pages = {
#                 0,
#                 total_pages // 2,
#                 total_pages - 1
#             }

#             for i in sample_pages:
#                 if i < 0 or i >= total_pages:
#                     continue

#                 page = pdf.pages[i]
#                 page_text = page.extract_text()

#                 if page_text:
#                     text += page_text + "\n"

#     except Exception as e:
#         print("PDF READ ERROR:", e)
#         return ""

#     if len(text.strip()) < 50:
#         print("⚠ Image-based PDF detected (OCR disabled on server)")
#         return ""

#     print(" Text PDF extracted")
#     return text[:8000]


# # ======================================
# # DOCX
# # ======================================
# def extract_docx_text(file_path):

#     try:
#         doc = docx.Document(file_path)
#         return "\n".join(p.text for p in doc.paragraphs)

#     except Exception as e:
#         print("DOCX ERROR:", e)
#         return ""


# # ======================================
# # TXT
# # ======================================
# def extract_txt_text(file_path):

#     try:
#         with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
#             return f.read()

#     except Exception as e:
#         print("TXT ERROR:", e)
#         return ""